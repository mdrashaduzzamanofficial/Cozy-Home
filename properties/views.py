from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Property, Application, Lease, Payment, UserPreference
from .forms import PropertyForm, ApplicationForm, LeaseForm, PaymentForm
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from docx import Document
from django.core.files.base import ContentFile
from io import BytesIO
import os
from django.conf import settings

def property_list(request):
    properties = Property.objects.filter(is_available=True)
    query = request.GET.get('q')
    if query:
        properties = properties.filter(location__icontains=query) | properties.filter(title__icontains=query)
    return render(request, 'properties/property_list.html', {'properties': properties})

def property_detail(request, pk):
    property = get_object_or_404(Property, pk=pk)
    if request.method == 'POST' and request.user.is_authenticated and request.user.role == 'tenant':
        form = ApplicationForm(request.POST)
        if form.is_valid():
            application = form.save(commit=False)
            application.tenant = request.user
            application.property = property
            application.save()
            messages.success(request, 'Application submitted successfully!')
            return redirect('property_detail', pk=pk)
    else:
        form = ApplicationForm()
    return render(request, 'properties/property_detail.html', {'property': property, 'form': form})

@login_required
def landlord_dashboard(request):
    if request.user.role != 'landlord':
        return redirect('property_list')
    properties = Property.objects.filter(owner=request.user)
    applications = Application.objects.filter(property__owner=request.user)
    return render(request, 'users/dashboard.html', {'properties': properties, 'applications': applications})

@login_required
def create_property(request):
    if request.user.role != 'landlord':
        return redirect('property_list')
    if request.method == 'POST':
        form = PropertyForm(request.POST, request.FILES)
        if form.is_valid():
            property = form.save(commit=False)
            property.owner = request.user
            property.save()
            messages.success(request, 'Property created successfully!')
            return redirect('landlord_dashboard')
    else:
        form = PropertyForm()
    return render(request, 'properties/property_form.html', {'form': form})

@login_required
def generate_lease(request, application_id):
    application = get_object_or_404(Application, pk=application_id, property__owner=request.user)
    if request.method == 'POST':
        form = LeaseForm(request.POST)
        if form.is_valid():
            lease = form.save(commit=False)
            lease.tenant = application.tenant
            lease.property = application.property
            doc = Document()
            doc.add_heading('Lease Agreement', 0)
            doc.add_paragraph(f"Tenant: {lease.tenant.username}")
            doc.add_paragraph(f"Property: {lease.property.title}")
            doc.add_paragraph(f"Start Date: {lease.start_date}")
            doc.add_paragraph(f"End Date: {lease.end_date}")
            doc.add_paragraph("Terms: This is a sample lease agreement.")
            buffer = BytesIO()
            doc.save(buffer)
            lease.file.save(f'lease_{lease.id}.docx', ContentFile(buffer.getvalue()))
            lease.save()
            application.status = 'approved'
            application.property.is_available = False
            application.save()
            application.property.save()
            messages.success(request, 'Lease generated successfully!')
            return redirect('landlord_dashboard')
    else:
        form = LeaseForm()
    return render(request, 'properties/lease_agreement.html', {'form': form, 'application': application})

@login_required
def make_payment(request, property_id):
    property = get_object_or_404(Property, pk=property_id)
    if request.method == 'POST':
        form = PaymentForm(request.POST)
        if form.is_valid():
            payment = form.save(commit=False)
            payment.tenant = request.user
            payment.property = property
            payment.status = 'completed'  # Simulated
            payment.save()
            messages.success(request, 'Payment recorded successfully!')
            return redirect('property_detail', pk=property.id)
    else:
        form = PaymentForm()
    return render(request, 'properties/payment_form.html', {'form': form, 'property': property})

def recommend_properties(request):
    if not request.user.is_authenticated or request.user.role != 'tenant':
        return redirect('property_list')
    try:
        preference = request.user.preference
        user_vector = np.array([[preference.max_budget, preference.min_bedrooms]])
        properties = Property.objects.filter(is_available=True)
        property_vectors = np.array([[p.price, p.bedrooms] for p in properties])
        if property_vectors.size == 0:
            return render(request, 'properties/property_list.html', {'properties': []})
        similarities = cosine_similarity(user_vector, property_vectors)
        sorted_properties = sorted(zip(properties, similarities[0]), key=lambda x: x[1], reverse=True)
        recommended = [p for p, _ in sorted_properties[:3]]  # Top 3
        return render(request, 'properties/property_list.html', {'properties': recommended})
    except UserPreference.DoesNotExist:
        messages.warning(request, 'Please set your preferences for recommendations.')
        return redirect('property_list')
    
"""
@login_required
def generate_lease(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    if request.user != application.property.owner:
        messages.error(request, "You are not authorized to generate a lease for this application.")
        return redirect('landlord_dashboard')

    if request.method == 'POST':
        form = LeaseForm(request.POST)
        if form.is_valid():
            lease = form.save(commit=False)
            lease.application = application
            lease.save()

            # Generate lease document
            doc = Document()
            doc.add_heading('Lease Agreement', 0)
            doc.add_paragraph(f'Tenant: {application.tenant.username}')
            doc.add_paragraph(f'Property: {application.property.title}')
            doc.add_paragraph(f'Start Date: {lease.start_date}')
            doc.add_paragraph(f'End Date: {lease.end_date}')
            doc.add_paragraph(f'Monthly Rent: ${lease.monthly_rent}')

            lease_filename = f'lease_{lease.id}.docx'
            lease_path = os.path.join(settings.MEDIA_ROOT, 'leases', lease_filename)
            doc.save(lease_path)

            lease.lease_document.name = os.path.join('leases', lease_filename)
            lease.save()

            messages.success(request, 'Lease generated successfully!')
            return redirect('landlord_dashboard')
    else:
        form = LeaseForm()
    return render(request, 'properties/lease_agreement.html', {'form': form, 'application': application}) 
    """